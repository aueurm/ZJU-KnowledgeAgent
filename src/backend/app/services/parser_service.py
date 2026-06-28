"""
PDF文件解析服务
职责：解析PDF/MD/TXT/DOCX教材文件，提取章节结构和内容
"""
import re
from pathlib import Path
from typing import List, Dict
from app.models import Chapter

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


class PDFParserService:
    """PDF解析服务"""

    def __init__(self):
        self.chapter_pattern = re.compile(r'^第[一二三四五六七八九十百千万\d]+章\s*')
        self.max_chars = 8000

    def parse(self, file_path: str) -> Dict:
        """解析PDF文件"""
        try:
            import fitz  # PyMuPDF
        except ImportError as exc:
            raise ImportError("PyMuPDF is not installed. Run: pip install -r requirements.txt") from exc

        doc = fitz.open(file_path)
        chapters = []
        current_chapter = None
        current_content = []
        chapter_index = 0
        all_content = []
        page_count = len(doc)

        for page_num, page in enumerate(doc):
            text = page.get_text()
            text = self._filter_header_footer(text, page_num)
            if text.strip():
                all_content.append(text)
            chapter_title = self._detect_chapter_title(text)

            if chapter_title:
                if current_chapter:
                    new_chapters = self._create_chapters(
                        current_chapter,
                        current_content,
                        page_num - len(current_content) + 1,
                        chapter_index
                    )
                    chapters.extend(new_chapters)
                    chapter_index += len(new_chapters)
                current_chapter = chapter_title
                current_content = [text]
            else:
                if current_chapter:
                    current_content.append(text)

        if current_chapter:
            new_chapters = self._create_chapters(
                current_chapter,
                current_content,
                page_num - len(current_content) + 2,
                chapter_index
            )
            chapters.extend(new_chapters)
        elif all_content:
            chapters.extend(self._create_chapters("全文", all_content, 1, 0))

        doc.close()
        return {"total_pages": page_count, "chapters": chapters}

    def _filter_header_footer(self, text: str, page_num: int) -> str:
        lines = text.split('\n')
        while lines and not lines[0].strip(): lines.pop(0)
        while lines and not lines[-1].strip(): lines.pop()
        return '\n'.join(lines)

    def _detect_chapter_title(self, text: str) -> str:
        lines = text.split('\n')
        for line in lines[:12]:
            line = line.strip()
            compact = re.sub(r'\s+', '', line)
            if line in ("绪论", "总论") or self.chapter_pattern.match(compact):
                return line
        return None

    def _create_chapters(self, title: str, contents: List[str], start_page: int, chapter_index: int) -> List[Chapter]:
        chapters = []
        current = []
        current_len = 0
        current_start = start_page

        for offset, text in enumerate(contents):
            if current and current_len + len(text) > self.max_chars:
                chapters.append(self._create_chapter(title, current, current_start, chapter_index + len(chapters)))
                current = []
                current_len = 0
                current_start = start_page + offset
            current.append(text)
            current_len += len(text)

        if current:
            chapters.append(self._create_chapter(title, current, current_start, chapter_index + len(chapters)))

        return chapters

    def _create_chapter(self, title: str, contents: List[str], start_page: int, chapter_index: int) -> Chapter:
        content = '\n'.join(contents)
        if len(content) > self.max_chars:
            content = content[:self.max_chars]
        return Chapter(
            chapter_id=f"ch_{chapter_index:03d}",
            title=title,
            page_start=start_page,
            page_end=start_page + len(contents) - 1,
            content=content,
            char_count=len(content)
        )


class MarkdownParserService:
    """Markdown解析服务"""

    def parse(self, file_path: str) -> Dict:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 按 # 标题分割章节
        sections = re.split(r'\n(?=#)', content)
        chapters = []
        chapter_id = 0

        for section in sections:
            lines = section.strip().split('\n')
            if not lines:
                continue

            # 第一行是标题
            title_line = lines[0].strip()
            if title_line.startswith('#'):
                title = title_line.lstrip('#').strip()
            else:
                title = f"第{chapter_id + 1}章"

            body = '\n'.join(lines[1:]).strip()
            if body:
                chapters.append(Chapter(
                    chapter_id=f"ch_{chapter_id:03d}",
                    title=title,
                    page_start=0,
                    page_end=0,
                    content=body,
                    char_count=len(body)
                ))
                chapter_id += 1

        return {"total_pages": 1, "chapters": chapters}


class TxtParserService:
    """TXT解析服务"""

    def parse(self, file_path: str) -> Dict:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 按空行分割段落，每1000字一组
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        chapters = []
        current_content = ""
        chapter_id = 0

        for para in paragraphs:
            if len(current_content) + len(para) > 1000 and current_content:
                chapters.append(Chapter(
                    chapter_id=f"ch_{chapter_id:03d}",
                    title=f"第{chapter_id + 1}节",
                    page_start=0,
                    page_end=0,
                    content=current_content,
                    char_count=len(current_content)
                ))
                chapter_id += 1
                current_content = para
            else:
                current_content += "\n\n" + para if current_content else para

        if current_content:
            chapters.append(Chapter(
                chapter_id=f"ch_{chapter_id:03d}",
                title=f"第{chapter_id + 1}节",
                page_start=0,
                page_end=0,
                content=current_content,
                char_count=len(current_content)
            ))

        return {"total_pages": 1, "chapters": chapters}


class DocxParserService:
    """DOCX解析服务"""

    def __init__(self):
        self.chapter_pattern = re.compile(r'^第[一二三四五六七八九十百千万\d]+章\s*')

    def parse(self, file_path: str) -> Dict:
        """解析DOCX文件"""
        if DocxDocument is None:
            raise ImportError("python-docx is not installed. Run: pip install python-docx")

        doc = DocxDocument(file_path)
        chapters = []
        current_title = None
        current_content = []
        chapter_id = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检测章节标题：Heading样式 或 匹配章节正则
            is_heading = (
                para.style.name.startswith('Heading')
                or self.chapter_pattern.match(text)
            )

            if is_heading:
                # 保存前一章
                if current_title and current_content:
                    body = '\n'.join(current_content)
                    chapters.append(Chapter(
                        chapter_id=f"ch_{chapter_id:03d}",
                        title=current_title,
                        page_start=0,
                        page_end=0,
                        content=body,
                        char_count=len(body)
                    ))
                    chapter_id += 1
                current_title = text
                current_content = []
            else:
                if current_title:
                    current_content.append(text)
                else:
                    # 标题前的内容归入"前言"
                    if not current_title:
                        current_title = "前言"
                    current_content.append(text)

        # 保存最后一章
        if current_title and current_content:
            body = '\n'.join(current_content)
            chapters.append(Chapter(
                chapter_id=f"ch_{chapter_id:03d}",
                title=current_title,
                page_start=0,
                page_end=0,
                content=body,
                char_count=len(body)
            ))

        # 如果没有检测到任何章节，将全文作为一个章节
        if not chapters:
            full_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            chapters.append(Chapter(
                chapter_id="ch_000",
                title="全文",
                page_start=0,
                page_end=0,
                content=full_text,
                char_count=len(full_text)
            ))

        return {"total_pages": len(doc.paragraphs) // 30 or 1, "chapters": chapters}


def get_parser_service(file_path: str):
    """根据文件扩展名返回对应的解析服务"""
    ext = Path(file_path).suffix.lower()
    if ext == '.pdf':
        return PDFParserService()
    elif ext == '.md':
        return MarkdownParserService()
    elif ext == '.txt':
        return TxtParserService()
    elif ext == '.docx':
        return DocxParserService()
    else:
        raise ValueError(f"Unsupported file format: {ext}")


# 向后兼容别名
MD_parser_service = MarkdownParserService
TXT_parser_service = TxtParserService
